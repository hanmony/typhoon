#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""D3/D4 的虚构边界测试；不读取真实资料或敏感源文件。"""
import contextlib
import io
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

import clean_docs_text as d4
import extract_docs as d3


class CleanRuleTests(unittest.TestCase):
    def test_r3_short_line_keeps_one(self):
        text = "页眉\n第一段。\n页眉\n第二段。\n页眉\n第三段。"
        cleaned, stats = d4.clean_text(text)
        self.assertEqual(stats["r3_dup_short"], 2)
        self.assertEqual(cleaned.count("页眉"), 1)

    def test_r3_long_line_keeps_one(self):
        header = "这是一个足够长的重复页眉用于测试清洗规则"
        text = "\n".join(f"{header}\n第{i}段正文。" for i in range(5))
        cleaned, stats = d4.clean_text(text)
        self.assertEqual(stats["r3_dup_long"], 4)
        self.assertEqual(cleaned.count(header), 1)

    def test_r3_numbered_list_repetitions_are_all_kept(self):
        item = "1、这是一个足够长且需要保留语境的合法正文列表项"
        text = "\n".join(f"{item}\n第{i}段结束。" for i in range(5))
        cleaned, stats = d4.clean_text(text)
        self.assertEqual(stats["r3_dup_long"], 0)
        self.assertEqual(cleaned.count(item), 5)

    def test_r2_adjacent_header_uses_pre_drop_snapshot(self):
        header = "这是一个足够长的重复页眉用于测试相邻页码删除"
        text = "\n".join(f"{header}\n{100 + i}" for i in range(5))
        _, stats = d4.clean_text(text)
        self.assertEqual(stats["r3_dup_long"], 4)
        self.assertEqual(stats["r2_adjacent"], 5)

    def test_r2_does_not_treat_numbered_body_as_header(self):
        item = "1、这是一个足够长且需要保留语境的合法正文列表项"
        text = "\n".join(f"{item}\n{100 + i}" for i in range(5))
        _, stats = d4.clean_text(text)
        self.assertEqual(stats["r2_adjacent"], 0)

    def test_r4_cap_never_truncates_content(self):
        first = "a" * 995
        second = "b" * 10
        merged, count, _ = d4.merge_lines([first, second])
        self.assertEqual(merged, [first, second])
        self.assertEqual(count, 0)
        self.assertEqual(sum(map(len, merged)), len(first) + len(second))

    def test_r4_does_not_cross_table_or_sheet_boundaries(self):
        self.assertFalse(d4.should_merge("普通正文", "单元格1 | 单元格2"))
        self.assertFalse(d4.should_merge("单元格1 | 单元格2", "普通正文"))
        self.assertFalse(d4.should_merge("普通正文", "=== sheet 1 ==="))
        self.assertFalse(d4.should_merge("=== sheet 1 ===", "普通正文"))

    def test_blank_lines_belong_to_r5_not_r3(self):
        cleaned, stats = d4.clean_text("正文一。\n\n\n\n正文二。")
        self.assertEqual(stats["r3_dup_short"], 0)
        self.assertEqual(cleaned, "正文一。\n\n正文二。")


class PipelineBoundaryTests(unittest.TestCase):
    def run_d4(self, meta, in_dir, out_dir):
        argv = ["clean_docs_text.py", str(meta), str(in_dir), str(out_dir)]
        output = io.StringIO()
        with mock.patch.object(sys, "argv", argv), contextlib.redirect_stdout(output):
            d4.main()
        return output.getvalue()

    def test_d4_removes_only_stale_txt_outputs(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            in_dir = root / "text"
            out_dir = root / "text_clean"
            in_dir.mkdir()
            out_dir.mkdir()
            (in_dir / "kept.txt").write_text("保留正文。", encoding="utf-8")
            (out_dir / "stale.txt").write_text("旧扫描产物", encoding="utf-8")
            (out_dir / "do-not-touch.bin").write_bytes(b"keep")
            meta = root / "extract_metadata.json"
            meta.write_text(json.dumps({"documents": [
                {"relpath": "kept.pdf", "status": "ok"}
            ]}, ensure_ascii=False), encoding="utf-8")

            self.run_d4(meta, in_dir, out_dir)

            self.assertTrue((out_dir / "kept.txt").is_file())
            self.assertFalse((out_dir / "stale.txt").exists())
            self.assertTrue((out_dir / "do-not-touch.bin").is_file())
            report = json.loads((root / "clean_report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["summary"]["stale_outputs_removed"], 1)

    def test_d4_rejects_sensitive_metadata_before_reading(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            in_dir = root / "text"
            out_dir = root / "text_clean"
            in_dir.mkdir()
            meta = root / "extract_metadata.json"
            meta.write_text(json.dumps({"documents": [
                {"relpath": "内部/值班表.pdf", "status": "ok"}
            ]}, ensure_ascii=False), encoding="utf-8")
            argv = ["clean_docs_text.py", str(meta), str(in_dir), str(out_dir)]
            output = io.StringIO()
            with mock.patch.object(sys, "argv", argv), contextlib.redirect_stdout(output):
                with self.assertRaises(SystemExit) as raised:
                    d4.main()
            self.assertEqual(raised.exception.code, 1)
            self.assertIn("敏感文件", output.getvalue())

    def test_scan_detection_ignores_repeated_lines_and_newlines(self):
        text = ("0 0 0 0\n" * 20) + ("\n" * 500) + "真实正文"
        self.assertEqual(d3.effective_body_chars(text), len("真实正文"))

    def test_docx_extracts_each_w_t_once_and_keeps_table_boundary(self):
        try:
            import docx
        except ImportError:
            self.skipTest("python-docx 未安装")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "synthetic.docx"
            document = docx.Document()
            document.add_paragraph("唯一段落")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "甲"
            table.cell(0, 1).text = "乙"
            document.save(path)

            text, error, extractor, note = d3.extract_docx(str(path))

            self.assertIsNone(error)
            self.assertIsNone(note)
            self.assertEqual(extractor, "python-docx")
            self.assertEqual(text.count("唯一段落"), 1)
            self.assertIn("甲 | 乙", text)


if __name__ == "__main__":
    unittest.main()
